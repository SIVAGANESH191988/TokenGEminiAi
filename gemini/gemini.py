from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import mysql.connector
from mysql.connector import Error
import os
import docx
from PyPDF2 import PdfReader
from PIL import Image, ImageEnhance
import pytesseract
import io
import re
from dotenv import load_dotenv
import google.generativeai as genai
from google.api_core.exceptions import ResourceExhausted
import extract_msg
from bs4 import BeautifulSoup
import json
import time
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
logger.info("Loading environment variables from .env file...")
load_dotenv()

# Set Tesseract path for Windows
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# MySQL Database Configuration
MYSQL_HOST = os.getenv("MYSQL_HOST", "localhost")
MYSQL_USER = os.getenv("MYSQL_USER", "root")
MYSQL_PASSWORD = os.getenv("MYSQL_PASSWORD", "Siyara@191988")
MYSQL_DATABASE = os.getenv("MYSQL_DATABASE", "geminidocument")

# Initialize FastAPI
app = FastAPI()

# Enable CORS to allow React frontend to communicate with backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize MySQL connection
def init_db():
    try:
        connection = mysql.connector.connect(
            host=MYSQL_HOST,
            user=MYSQL_USER,
            password=MYSQL_PASSWORD,
            database=MYSQL_DATABASE
        )
        if connection.is_connected():
            cursor = connection.cursor()
            create_table_query = """
            CREATE TABLE IF NOT EXISTS extracted_data (
                id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(255),
                email VARCHAR(255),
                number VARCHAR(50),
                professional_summary TEXT,
                project_name TEXT,
                skills TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
            """
            cursor.execute(create_table_query)
            connection.commit()
            return connection, cursor
    except Error as e:
        logger.error(f"Database connection failed: {e}")
        return None, None

# Load Gemini API Key
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
logger.info(f"GEMINI_API_KEY found: {'Yes' if GEMINI_API_KEY else 'No'}")
if not GEMINI_API_KEY:
    logger.error("GEMINI_API_KEY not found in .env file. Please add it to the .env file.")
    raise HTTPException(status_code=500, detail="GEMINI_API_KEY not found in .env file.")
genai.configure(api_key=GEMINI_API_KEY)

# Configure Gemini model
model = genai.GenerativeModel(
    model_name="models/gemini-1.5-flash",
    generation_config={
        "temperature": 0.0,
        "top_p": 1,
        "top_k": 1,
        "max_output_tokens": 1024,
    }
)

def count_tokens(text):
    return len(text) // 4 + 1 if text else 0

def read_file(file: UploadFile):
    filename = file.filename.lower()
    all_content = []
    try:
        content = file.file.read()
        file.file.seek(0)
        if filename.endswith(".msg"):
            msg = extract_msg.Message(io.BytesIO(content))
            if msg.htmlBody:
                soup = BeautifulSoup(msg.htmlBody, 'html.parser')
                all_content.append(soup.get_text(separator=' ', strip=True))
            elif msg.body:
                all_content.append(msg.body)
            if msg.attachments:
                for att in msg.attachments:
                    fname = re.sub(r'[^\w\.\-]', '', att.longFilename or att.shortFilename or "Unnamed")
                    att_content = io.BytesIO(att.data)
                    if fname.endswith(".pdf"):
                        reader = PdfReader(att_content)
                        content = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
                        if content.strip():
                            all_content.append(content)
                    elif fname.endswith(".docx"):
                        doc = docx.Document(att_content)
                        content = "\n".join(p.text for p in doc.paragraphs)
                        if content.strip():
                            all_content.append(content)
                    elif fname.endswith((".jpg", ".jpeg", ".png")):
                        image = Image.open(att_content)
                        image = ImageEnhance.Contrast(image).enhance(2.0)
                        content = pytesseract.image_to_string(image, lang='eng')
                        if content.strip():
                            all_content.append(content)
            return "\n".join(all_content)[:50000]
        elif filename.endswith(".txt"):
            return content.decode("utf-8")[:50000]
        elif filename.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())[:50000]
        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)[:50000]
        elif filename.endswith((".jpg", ".jpeg", ".png")):
            image = Image.open(io.BytesIO(content))
            image = ImageEnhance.Contrast(image).enhance(2.0)
            content = pytesseract.image_to_string(image, lang='eng')
            return content[:50000] if content.strip() else ""
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}")
    except Exception as e:
        logger.error(f"Error processing file {filename}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

def extract_data_with_gemini(content, retries=3, delay=10):
    if not content:
        return None, 0, 0
    prompt = f"""
The text below contains content from an email and documents. Extract structured data and return as a list of JSON objects with the following keys:
- name
- email
- number
- professional_summary
- project_name
- skills
Ensure all field values are strings (convert lists to comma-separated strings if needed).
Format:
[
    {{
        "name": "...",
        "email": "...",
        "number": "...",
        "professional_summary": "...",
        "project_name": "...",
        "skills": "..."
    }},
    ...
]
Text: {content}
"""
    input_tokens = count_tokens(prompt)
    for attempt in range(retries):
        try:
            response = model.generate_content(prompt)
            result = response.text.strip()
            result = re.sub(r'^json\n?|```json|```', '', result, flags=re.MULTILINE).strip()
            output_tokens = count_tokens(result)
            return result, input_tokens, output_tokens
        except ResourceExhausted:
            if attempt < retries - 1:
                time.sleep(delay)
                delay *= 2
            else:
                logger.error("Gemini API quota exhausted")
                raise HTTPException(status_code=429, detail="Gemini API quota exhausted")
        except Exception as e:
            logger.error(f"Gemini error: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Gemini error: {str(e)}")

def detect_email_intent(content):
    prompt = (
        "Read the email content below and identify the main intent or purpose in one line.\n"
        "Examples: Invitation for interview, Rejection letter, Offer letter, Request for documents, Acknowledgement, General update\n\n"
        f"Email Content:\n\"\"\"\n{content}\n\"\"\"\n\nRespond with only the intent"
    )
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        logger.warning(f"Intent detection failed: {str(e)}")
        return "Unknown"

def convert_lists_to_strings(record):
    return {k: ", ".join(v) if isinstance(v, list) else str(v) for k, v in record.items()}

def is_duplicate(record, connection, cursor):
    try:
        record = convert_lists_to_strings(record)
        query = """
        SELECT COUNT(*) FROM extracted_data 
        WHERE name = %s AND email = %s AND number = %s
        """
        values = (
            record.get("name", ""),
            record.get("email", ""),
            record.get("number", "")
        )
        cursor.execute(query, values)
        return cursor.fetchone()[0] > 0
    except Exception as e:
        logger.warning(f"Error checking duplicates: {str(e)}")
        return False

def store_data_in_db(data, connection, cursor):
    try:
        parsed = json.loads(data)
        inserted = False
        for record in parsed:
            record = convert_lists_to_strings(record)
            if is_duplicate(record, connection, cursor):
                logger.info(f"Skipped storing duplicate record for {record.get('name', 'Unknown')}")
                continue
            insert_query = """
            INSERT INTO extracted_data (name, email, number, professional_summary, project_name, skills)
            VALUES (%s, %s, %s, %s, %s, %s)
            """
            values = (
                record.get("name", ""),
                record.get("email", ""),
                record.get("number", ""),
                record.get("professional_summary", ""),
                record.get("project_name", ""),
                record.get("skills", "")
            )
            cursor.execute(insert_query, values)
            inserted = True
        if inserted:
            connection.commit()
        return inserted
    except json.JSONDecodeError:
        logger.error("Failed to parse JSON")
        raise HTTPException(status_code=400, detail="Failed to parse JSON")
    except Exception as e:
        logger.error(f"Database storage error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Database storage error: {str(e)}")

def fetch_ids(connection, cursor):
    try:
        cursor.execute("SELECT id FROM extracted_data")
        return [str(row[0]) for row in cursor.fetchall()]
    except Exception as e:
        logger.error(f"Error fetching IDs: {str(e)}")
        raise HTTPException(status_code=500, detail="Error fetching IDs")

def fetch_record_by_id(record_id, connection, cursor):
    try:
        query = """
        SELECT name, email, number, professional_summary, project_name, skills
        FROM extracted_data WHERE id = %s
        """
        cursor.execute(query, (record_id,))
        result = cursor.fetchone()
        if result:
            return {
                "name": result[0] or "N/A",
                "email": result[1] or "N/A",
                "number": result[2] or "N/A",
                "professional_summary": result[3] or "N/A",
                "project_name": result[4] or "N/A",
                "skills": result[5] or "N/A"
            }
        return None
    except Exception as e:
        logger.error(f"Error fetching record: {str(e)}")
        raise HTTPException(status_code=500, detail="Error fetching record")

def fetch_all_records(connection, cursor):
    try:
        query = """
        SELECT id, name, email, number, professional_summary, project_name, skills
        FROM extracted_data
        """
        cursor.execute(query)
        results = cursor.fetchall()
        records = []
        for result in results:
            records.append({
                "id": str(result[0]),
                "name": result[1] or "N/A",
                "email": result[2] or "N/A",
                "number": result[3] or "N/A",
                "professional_summary": result[4] or "N/A",
                "project_name": result[5] or "N/A",
                "skills": result[6] or "N/A"
            })
        return records
    except Exception as e:
        logger.error(f"Error fetching all records: {str(e)}")
        raise HTTPException(status_code=500, detail="Error fetching all records")

# API Endpoints
@app.post("/upload")
async def upload_files(files: list[UploadFile] = File(...)):
    connection, cursor = init_db()
    if not connection or not cursor:
        raise HTTPException(status_code=500, detail="Database connection failed")
    
    results = []
    total_tokens = 0
    try:
        for file in files:
            content = read_file(file)
            intent = detect_email_intent(content)
            result, in_tokens, out_tokens = extract_data_with_gemini(content)
            total_tokens += in_tokens + out_tokens
            if result:
                inserted = store_data_in_db(result, connection, cursor)
                results.append({
                    "filename": file.filename,
                    "intent": intent,
                    "extracted_data": json.loads(result) if result else None,
                    "tokens_used": in_tokens + out_tokens,
                    "stored": inserted
                })
            else:
                results.append({
                    "filename": file.filename,
                    "intent": intent,
                    "extracted_data": None,
                    "tokens_used": in_tokens,
                    "stored": False
                })
        return {"results": results, "total_tokens": total_tokens}
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()

@app.get("/records")
async def get_records():
    connection, cursor = init_db()
    if not connection or not cursor:
        raise HTTPException(status_code=500, detail="Database connection failed")
    try:
        ids = fetch_ids(connection, cursor)
        return {"ids": ids}
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()

@app.get("/record/{record_id}")
async def get_record(record_id: str):
    connection, cursor = init_db()
    if not connection or not cursor:
        raise HTTPException(status_code=500, detail="Database connection failed")
    try:
        record = fetch_record_by_id(record_id, connection, cursor)
        if not record:
            raise HTTPException(status_code=404, detail="Record not found")
        return record
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()

@app.get("/all_records")
async def get_all_records():
    connection, cursor = init_db()
    if not connection or not cursor:
        raise HTTPException(status_code=500, detail="Database connection failed")
    try:
        records = fetch_all_records(connection, cursor)
        return {"records": records}
    finally:
        if connection.is_connected():
            cursor.close()
            connection.close()